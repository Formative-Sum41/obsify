"""Generate a SYNTHETIC complex corpus for testing obsify end to end / as a demo.

Everything here is FAKE. ABN/ACN/TFN are checksum-valid (minted with obsify's own
validators) so the recognizers exercise their real logic. The corpus spans all three
supported formats and the hard cases:

  * multi-sheet Excel — a numeric-heavy journal (sequential IDs + alnum account codes
    that must NOT be flagged as TFN/ACN/PERSON), plus vendor/employee sheets with
    inline-labelled ABN/TFN, names and emails;
  * a PDF engagement letter — prose (names, AU address, email, engagement code, labelled
    TFN) plus a trial-balance table with the client name + ABN/ACN;
  * a DOCX audit memo — paragraphs + a vendor table (labelled ABN).

Excel and DOCX generation always work. PDF generation needs reportlab
(`pip install "obsify[demo]"`); without it the PDF is skipped with a note.

Run:  python -m obsify.make_corpus --out ./corpus_demo
"""

from __future__ import annotations

import argparse
from pathlib import Path

from obsify.checksums import complete_acn, is_valid_abn, is_valid_tfn


# --- checksum-valid synthetic identifiers ------------------------------------

def _mint_abn(base: int) -> str:
    n = base
    while not is_valid_abn(str(n)):
        n += 1
    return str(n)


def _mint_tfn(base: int) -> str:
    n = base
    while not is_valid_tfn(str(n)):
        n += 1
    return str(n)


ABN = _mint_abn(41628054000)          # Marranbah
ACN = complete_acn("41628054")
VENDOR_ABN = _mint_abn(72155399000)   # Halloway Supplies
TFN = _mint_tfn(284551900)


def _fmt_abn(a: str) -> str:
    return f"{a[:2]} {a[2:5]} {a[5:8]} {a[8:]}"


def _fmt_acn(a: str) -> str:
    return f"{a[:3]} {a[3:6]} {a[6:]}"


def _fmt_tfn(a: str) -> str:
    return f"{a[:3]} {a[3:6]} {a[6:]}"


# --- planted (all synthetic) — also the answer key a test can assert against --

CLIENTS = ["Marranbah Pastoral Co Pty Ltd", "Kellview Logistics Pty Ltd",
           "Ashgrove Nominees Pty Ltd"]
PEOPLE = ["Elena Marsh", "Tomas Vidal", "Nadia Okonkwo", "Grigor Sallo"]
EMAILS = ["e.marsh@example.com", "t.vidal@kellview.example.com"]
ADDRESS = "Unit 7, 14 Sturt Street, Ballarat VIC 3350"
ENGAGEMENT_CODE = "AUD-771204"


def planted() -> dict:
    """The synthetic PII planted in the corpus, by type — an answer key for tests."""
    return {
        "clients": CLIENTS,
        "AU_ABN": [_fmt_abn(ABN), _fmt_abn(VENDOR_ABN)],
        "AU_ACN": [_fmt_acn(ACN)],
        "AU_TFN": [_fmt_tfn(TFN)],
        "PERSON": PEOPLE,
        "EMAIL_ADDRESS": EMAILS,
        "ADDRESS": [ADDRESS],
        "ENGAGEMENT_CODE": [ENGAGEMENT_CODE],
    }


# --- Excel -------------------------------------------------------------------

def build_excel(path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()

    # Sheet 1: a numeric-heavy journal — the false-positive minefield.
    gl = wb.active
    gl.title = "Journals"
    gl.append(["JournalID", "Date", "Account", "Debit", "Credit", "Description"])
    # 8-digit sequential JournalIDs: some will pass the TFN/ACN checksum by chance,
    # yet must NOT be flagged — they carry no "TFN"/"ACN" context word. This is the
    # exact false-positive that floods real numeric ledgers.
    rows = [
        (18400001, "2025-06-30", "EP180102", 225140.00, 0, "Cash at bank reconciliation"),
        (18400002, "2025-06-30", "EP180233", 0, 51000.00, f"Distribution to {CLIENTS[2]}"),
        (18400003, "2025-06-29", "GJ180614", 88000.00, 0,
         f"Term deposit for {CLIENTS[0]} rolled over"),
        (18400004, "2025-06-29", "EP180440", 12750.50, 0, "Freight accrual"),
        (18400005, "2025-06-28", "GJ180705", 0, 4180.00,
         "Reimbursement; in-cell wrap Kellview_x000A_Logistics Pty Ltd"),
        (18400006, "2025-06-28", "EP180512", 903312.00, 0, "Intercompany settlement"),
        (18400007, "2025-06-27", "GJ180801", 0, 1890.00, "Sundry expense"),
    ]
    for r in rows:
        gl.append(list(r))

    # Sheet 2: vendors — ABN labelled INLINE (per-cell segments never see a header
    # in another cell, so context must live in the same cell to be detected).
    ven = wb.create_sheet("Vendors")
    ven.append(["Vendor", "Registration", "Contact", "Email"])
    ven.append(["Halloway Supplies Pty Ltd", f"ABN {_fmt_abn(VENDOR_ABN)}",
                PEOPLE[3], "accounts@halloway.example.com"])
    ven.append([CLIENTS[0], f"ABN {_fmt_abn(ABN)} / ACN {_fmt_acn(ACN)}",
                PEOPLE[1], EMAILS[1]])

    # Sheet 3: employees — TFN labelled inline, names, emails.
    emp = wb.create_sheet("Employees")
    emp.append(["Name", "Role", "TaxNote", "Email"])
    emp.append([PEOPLE[0], "Engagement partner", f"TFN {_fmt_tfn(TFN)}", EMAILS[0]])
    emp.append([PEOPLE[2], "Analyst", "TFN pending", "n.okonkwo@example.com"])

    wb.save(str(path))


# --- PDF (optional: needs reportlab) -----------------------------------------

def build_pdf(path: Path) -> bool:
    """Render the engagement-letter PDF. Returns False (with no file) if reportlab
    is not installed — the rest of the corpus is unaffected."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                        TableStyle)
    except Exception:
        return False

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
                            topMargin=20 * mm, bottomMargin=20 * mm)
    story = [
        Paragraph("Engagement Letter - Statutory Audit 2025", styles["Title"]),
        Spacer(1, 4 * mm),
        Paragraph(f"Dear Mr {PEOPLE[1]},", styles["BodyText"]),
        Paragraph(f"We confirm our appointment for {CLIENTS[0]} (reference "
                  f"{ENGAGEMENT_CODE}), led by partner {PEOPLE[0]}, who can be reached at "
                  f"{EMAILS[0]}.", styles["BodyText"]),
        Paragraph(f"Please send documents to {ADDRESS}. Your nominated tax file number "
                  f"(TFN {_fmt_tfn(TFN)}) is handled per our privacy obligations.",
                  styles["BodyText"]),
        Spacer(1, 6 * mm),
        Paragraph("Trial Balance Extract (AUD)", styles["Heading3"]),
    ]
    data = [["Account", "Description", "Debit", "Credit"],
            ["1000", "Cash at bank", "225,140.00", ""],
            ["1200", "Trade receivables", "410,662.00", ""],
            ["3000", "Retained earnings", "", "402,118.00"]]
    tbl = Table(data, colWidths=[24 * mm, 70 * mm, 34 * mm, 34 * mm])
    tbl.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                             ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                             ("FONTSIZE", (0, 0), (-1, -1), 9)]))
    story += [tbl, Spacer(1, 6 * mm),
              Paragraph(f"1. {CLIENTS[0]} (ABN {_fmt_abn(ABN)}; ACN {_fmt_acn(ACN)}).",
                        styles["BodyText"])]
    doc.build(story)
    return True


# --- DOCX --------------------------------------------------------------------

def build_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Audit Planning Memo - 2025", level=1)
    doc.add_paragraph(f"Client: {CLIENTS[1]}. Prepared by {PEOPLE[0]} "
                      f"(engagement {ENGAGEMENT_CODE}).")
    doc.add_paragraph(f"Primary contact {PEOPLE[1]} confirmed the registered office at "
                      f"{ADDRESS}, reachable at {EMAILS[1]}.")
    doc.add_paragraph(f"Personnel note: analyst {PEOPLE[2]}; partner tax reference "
                      f"TFN {_fmt_tfn(TFN)} on file.")

    doc.add_heading("Key vendors", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Vendor", "Registration", "Contact"
    r = table.add_row().cells
    r[0].text, r[1].text, r[2].text = ("Halloway Supplies Pty Ltd",
                                       f"ABN {_fmt_abn(VENDOR_ABN)}", PEOPLE[3])
    doc.save(str(path))


# --- driver ------------------------------------------------------------------

def build_corpus(out_dir: str) -> dict:
    """Write the corpus to out_dir. Returns {'files': [...], 'planted': {...},
    'notes': [...]}. Excel + DOCX always written; PDF only if reportlab is present."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    files: list[str] = []
    notes: list[str] = []

    xlsx = out / "ledger.xlsx"
    build_excel(xlsx)
    files.append(str(xlsx))

    docx = out / "audit_memo.docx"
    build_docx(docx)
    files.append(str(docx))

    pdf = out / "engagement_letter.pdf"
    if build_pdf(pdf):
        files.append(str(pdf))
    else:
        notes.append("reportlab not installed; skipped engagement_letter.pdf "
                     "(install with: pip install \"obsify[demo]\")")

    # A master list (client names) — handy for demonstrating differential tooling.
    master = out / "master_list.txt"
    master.write_text("\n".join(CLIENTS) + "\n", encoding="utf-8")
    files.append(str(master))

    return {"files": files, "planted": planted(), "notes": notes}


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(prog="python -m obsify.make_corpus",
                                 description="Generate a synthetic complex corpus (all fake).")
    ap.add_argument("--out", default="./corpus_demo", help="output directory")
    args = ap.parse_args(argv)
    result = build_corpus(args.out)
    print(f"[make_corpus] wrote {len(result['files'])} file(s) to {args.out}:")
    for f in result["files"]:
        print(f"  - {Path(f).name}")
    for n in result["notes"]:
        print(f"[make_corpus] note: {n}")
    print("[make_corpus] all values are synthetic; ABN/ACN/TFN are checksum-valid.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
