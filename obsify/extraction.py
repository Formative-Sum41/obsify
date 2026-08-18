"""Text extraction from PDF, Excel and Word (.docx) documents.

Produces a flat list of `Segment`s — each a normalized plain-text span tagged with a
precise source location (page/table/cell) so a detection can be traced back to where
it lives. Extraction performs no detection and no judgement.

Extractors are deterministic and read-only. camelot is a lazy, optional fallback for
complex tables (it needs a system Ghostscript binary); when unavailable, obsify records
a note and proceeds on pdfplumber output rather than failing — graceful degradation
with an explicit signal.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import openpyxl
import pdfplumber

_WS = re.compile(r"\s+")
# OOXML escapes control characters in shared strings as _xXXXX_. openpyxl does
# not always unescape these, so an in-cell newline can surface as the LITERAL
# text "_x000A_" (LF), "_x000D_" (CR), "_x0009_" (tab). A plain "\s+" pass never
# matches that literal, so the break survives into the segment and defeats
# matching (the exact "Kestrel Bay\nHoldings" extraction miss). Fold every
# OOXML-escaped control character (0x00-0x1F) to a space first.
_XML_CTRL = re.compile(r"_x00[01][0-9A-Fa-f]_")


def normalize_ws(text: str) -> str:
    """Collapse all whitespace runs (incl. newlines/tabs and OOXML-escaped line
    breaks) to single spaces.

    This is what joins a value split across a line break *within* a cell or
    paragraph into a single matchable span.
    """
    text = _XML_CTRL.sub(" ", text)
    return _WS.sub(" ", text.replace(" ", " ")).strip()


@dataclass(frozen=True)
class Segment:
    document: str  # source file name (basename only; no path leakage into logs)
    locator: str   # human-readable source location, e.g. "page 2 / table 1 / r3c2"
    kind: str      # prose | table_cell | excel_cell | excel_header
    text: str      # normalized plain text


@dataclass
class PageCoverage:
    document: str
    page: int
    chars: int
    median: int
    flag: str  # "" if fine, else a LOW COVERAGE message


# A page extracting fewer than this fraction of the document's median character
# count (or zero) is flagged as possible under-extraction / image content.
LOW_COVERAGE_FRACTION = 0.20


# --- camelot lazy fallback ---------------------------------------------------

def _try_camelot(pdf_path: str, page_number: int, notes: list[str]) -> list[Segment] | None:
    """Attempt a camelot table extraction for one page. Returns segments or None.

    Imported lazily because camelot pulls in Ghostscript/opencv; a missing
    dependency must degrade gracefully, not crash the run.
    """
    try:
        import camelot  # noqa: PLC0415  (intentional lazy import)
    except Exception as exc:  # ImportError or downstream dependency failure
        note = f"camelot unavailable ({type(exc).__name__}); relied on pdfplumber only"
        if note not in notes:
            notes.append(note)
        return None

    document = Path(pdf_path).name
    segments: list[Segment] = []
    try:
        tables = camelot.read_pdf(pdf_path, pages=str(page_number), flavor="lattice")
        for t_idx, table in enumerate(tables):
            for r_idx, row in enumerate(table.df.values.tolist()):
                for c_idx, cell in enumerate(row):
                    norm = normalize_ws(str(cell))
                    if norm:
                        segments.append(Segment(
                            document=document,
                            locator=f"page {page_number} / camelot-table {t_idx + 1} / r{r_idx + 1}c{c_idx + 1}",
                            kind="table_cell",
                            text=norm,
                        ))
    except Exception as exc:
        notes.append(f"camelot failed on {document} page {page_number}: {type(exc).__name__}")
        return None
    return segments or None


# --- PDF ---------------------------------------------------------------------

def extract_pdf(pdf_path: str, notes: list[str]) -> tuple[list[Segment], list[int]]:
    """Return (segments, chars_per_page). chars_per_page drives the coverage check."""
    document = Path(pdf_path).name
    segments: list[Segment] = []
    page_chars: list[int] = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_no = page_index + 1
            chars = 0

            # Tables first (structured cells give the most precise locators).
            tables = page.extract_tables() or []
            for t_idx, table in enumerate(tables):
                for r_idx, row in enumerate(table):
                    for c_idx, cell in enumerate(row):
                        if cell is None:
                            continue
                        norm = normalize_ws(cell)
                        if norm:
                            chars += len(norm)
                            segments.append(Segment(
                                document=document,
                                locator=f"page {page_no} / table {t_idx + 1} / r{r_idx + 1}c{c_idx + 1}",
                                kind="table_cell",
                                text=norm,
                            ))

            # If pdfplumber found no tables, try camelot as a fallback.
            if not tables:
                camelot_segs = _try_camelot(pdf_path, page_no, notes)
                if camelot_segs:
                    for seg in camelot_segs:
                        chars += len(seg.text)
                    segments.extend(camelot_segs)

            # Prose: line-level segments give useful locators for footnotes.
            text = page.extract_text() or ""
            for line_no, raw_line in enumerate(text.splitlines(), start=1):
                norm = normalize_ws(raw_line)
                if norm:
                    chars += len(norm)
                    segments.append(Segment(
                        document=document,
                        locator=f"page {page_no} / line {line_no}",
                        kind="prose",
                        text=norm,
                    ))

            # A page-level block joins all lines into one whitespace-normalized
            # span, so entities that wrap across visual lines (long addresses,
            # multi-word client names) are still seen as one contiguous string.
            # Not added to the character count (it duplicates the line text).
            block = normalize_ws(text)
            if block:
                segments.append(Segment(
                    document=document,
                    locator=f"page {page_no} / full text",
                    kind="prose_block",
                    text=block,
                ))

            page_chars.append(chars)

    return segments, page_chars


# --- Excel -------------------------------------------------------------------

def extract_excel(xlsx_path: str) -> list[Segment]:
    document = Path(xlsx_path).name
    segments: list[Segment] = []

    # data_only=True yields computed values rather than formulae — obsify scans the
    # visible content an analyst would see.
    wb = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
    try:
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    norm = normalize_ws(str(cell.value))
                    if not norm:
                        continue
                    kind = "excel_header" if cell.row == 1 else "excel_cell"
                    segments.append(Segment(
                        document=document,
                        locator=f"sheet '{sheet.title}' / cell {cell.coordinate}",
                        kind=kind,
                        text=norm,
                    ))
    finally:
        wb.close()
    return segments


# --- DOCX --------------------------------------------------------------------

def extract_docx(docx_path: str, notes: list[str]) -> list[Segment]:
    """Extract paragraphs and table cells from a .docx as located Segments.

    python-docx is imported lazily so a missing dependency degrades gracefully
    (note + skip) rather than crashing a multi-file run — same posture as camelot.
    Each paragraph's runs are already joined by python-docx, so a value split
    across runs surfaces as one contiguous span.
    """
    try:
        from docx import Document  # noqa: PLC0415  (lazy: optional at runtime)
    except Exception as exc:  # ImportError or downstream failure
        notes.append(f"python-docx unavailable ({type(exc).__name__}); "
                     f"skipped {Path(docx_path).name} (unread blind spot)")
        return []

    document = Path(docx_path).name
    segments: list[Segment] = []
    doc = Document(docx_path)

    for p_idx, para in enumerate(doc.paragraphs, start=1):
        norm = normalize_ws(para.text)
        if norm:
            segments.append(Segment(document=document, locator=f"paragraph {p_idx}",
                                    kind="prose", text=norm))

    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                norm = normalize_ws(cell.text)
                if norm:
                    segments.append(Segment(
                        document=document,
                        locator=f"table {t_idx} / r{r_idx}c{c_idx}",
                        kind="table_cell", text=norm))
    return segments


# --- dispatch ----------------------------------------------------------------

_PDF_SUFFIXES = {".pdf"}
_EXCEL_SUFFIXES = {".xlsx", ".xlsm"}
_DOCX_SUFFIXES = {".docx"}
SUPPORTED_SUFFIXES = _PDF_SUFFIXES | _EXCEL_SUFFIXES | _DOCX_SUFFIXES


def _median(values: list[int]) -> int:
    if not values:
        return 0
    s = sorted(values)
    n = len(s)
    mid = n // 2
    return s[mid] if n % 2 else (s[mid - 1] + s[mid]) // 2


def _page_coverage(document: str, page_chars: list[int]) -> list[PageCoverage]:
    """Flag pages that extracted no text, or far less than the document's median."""
    median = _median(page_chars)
    threshold = median * LOW_COVERAGE_FRACTION
    coverage: list[PageCoverage] = []
    for i, chars in enumerate(page_chars):
        if chars == 0:
            flag = "LOW COVERAGE - empty page (no text layer / image content)"
        elif median > 0 and chars < threshold:
            flag = "LOW COVERAGE - sparse relative to document median"
        else:
            flag = ""
        coverage.append(PageCoverage(document=document, page=i + 1, chars=chars,
                                     median=median, flag=flag))
    return coverage


def extract_document(path: str, notes: list[str]) -> tuple[list[Segment], list[PageCoverage]]:
    suffix = Path(path).suffix.lower()
    document = Path(path).name
    if suffix not in SUPPORTED_SUFFIXES:
        # Unknown types are skipped with an explicit note rather than silently.
        notes.append(f"skipped unsupported file type: {document} ({suffix or 'no suffix'})")
        return [], []
    # A corrupt/encrypted/old-format file (e.g. real .xls saved as .xlsx, a
    # password-protected book) must be skipped with a note, never crash the run —
    # an unread file is a blind spot to surface, not a halted pipeline.
    try:
        if suffix in _PDF_SUFFIXES:
            segments, page_chars = extract_pdf(path, notes)
            return segments, _page_coverage(document, page_chars)
        if suffix in _DOCX_SUFFIXES:
            return extract_docx(path, notes), []  # DOCX is not paged; no coverage check
        return extract_excel(path), []  # Excel is not paged; no coverage check
    except Exception as exc:
        notes.append(f"could not read {document}: {type(exc).__name__}; SKIPPED (unread blind spot)")
        return [], []
