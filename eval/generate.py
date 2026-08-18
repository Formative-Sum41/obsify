"""Generate a LABELLED evaluation corpus (all synthetic) + a ground-truth answer key.

This is dev/eval tooling — it lives outside the shipped `obsify` package and is not in
the wheel. It exists to produce measurable precision/recall numbers, so every planted
item is tagged with its EXPECTED behaviour under obsify's shipping detector:

  expected = "detect"    -> should be flagged (in-context ID, name, email, card, ...)
  expected = "suppress"  -> should be DROPPED by policy (a bare ABN/ACN/TFN with no
                            nearby label word — require-context is a feature, not a miss)
  expected = "gap"       -> obsify has no recognizer for this type (Medicare / IP / DOB);
                            0% detection is a documented coverage gap, not a bug

Difficulty is concentrated where it is informative: adversarial ID formats, values split
by in-cell newlines, unicode names, missing context, and FP-TORTURE numeric columns
(sequential 8-9 digit IDs, amounts, alnum codes) that plant NO PII — any detection there
is an emergent false positive, the most honest precision signal.

Run:  python eval/generate.py --out ./eval_corpus
Emits: eval_corpus/*.xlsx|pdf|docx  and  eval_corpus/answer_key.jsonl
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from faker import Faker

from obsify.checksums import complete_acn, complete_medicare, is_valid_abn, is_valid_tfn

fake = Faker("en_AU")
Faker.seed(7)

# --------------------------------------------------------------- id minting ----

def _mint_abn(base: int) -> str:
    while not is_valid_abn(str(base)):
        base += 1
    return str(base)


def _mint_tfn(base: int) -> str:
    while not is_valid_tfn(str(base)):
        base += 1
    return str(base)


def fmt_abn(a: str) -> str:
    return f"{a[:2]} {a[2:5]} {a[5:8]} {a[8:]}"


def fmt_acn(a: str) -> str:
    return f"{a[:3]} {a[3:6]} {a[6:]}"


def fmt_tfn(a: str) -> str:
    return f"{a[:3]} {a[3:6]} {a[6:]}"


ABN1, ABN2 = _mint_abn(51824753000), _mint_abn(29847115000)
ACN1 = complete_acn("29119004")
TFN1, TFN2, TFN3 = _mint_tfn(310221400), _mint_tfn(555444300), _mint_tfn(863291700)

# --------------------------------------------------------------- answer key ----

_KEY: list[dict] = []


def plant(document: str, type_: str, value: str, expected: str, note: str = "") -> str:
    """Record a planted item and return its value (so it can be embedded verbatim)."""
    _KEY.append({"document": document, "type": type_, "value": value,
                 "expected": expected, "note": note})
    return value


# --------------------------------------------------------------- Excel ---------

def build_excel(path: Path) -> None:
    import openpyxl
    doc = path.name
    wb = openpyxl.Workbook()

    # Sheet 1 — Contacts: names, emails, phones, in-context vs bare TFN, unicode.
    ws = wb.active
    ws.title = "Contacts"
    ws.append(["Name", "Email", "Phone", "TaxNote", "TaxIdBare"])
    p1 = plant(doc, "PERSON", "Priyanka Raman", "detect", "clean name")
    p2 = plant(doc, "PERSON", "José Márquez", "detect", "unicode name")
    e1 = plant(doc, "EMAIL_ADDRESS", "p.raman@example.com", "detect", "email")
    e2 = plant(doc, "EMAIL_ADDRESS", "jose.marquez@example.com", "detect", "email")
    tfn_inline = plant(doc, "AU_TFN", fmt_tfn(TFN1), "detect", "inline-labelled TFN")
    tfn_bare = plant(doc, "AU_TFN", fmt_tfn(TFN2), "suppress", "bare TFN, no context word")
    ws.append([p1, e1, "0412 345 678", f"TFN {tfn_inline}", tfn_bare])
    ws.append([p2, e2, "(02) 5550 1234", "TFN pending", ""])
    # in-cell newline splitting a name (must still match after normalize_ws)
    p3 = plant(doc, "PERSON", "Nadia Okonkwo", "detect", "name split by in-cell newline")
    ws.append(["Nadia_x000A_Okonkwo", "n.okonkwo@example.com", "0455 662 118", "TFN pending", ""])
    plant(doc, "EMAIL_ADDRESS", "n.okonkwo@example.com", "detect", "email")

    # Sheet 2 — Vendors: inline vs bare ABN, ACN, credit card, IBAN.
    vs = wb.create_sheet("Vendors")
    vs.append(["Vendor", "Registration", "Card", "IBAN"])
    abn_inline = plant(doc, "AU_ABN", fmt_abn(ABN1), "detect", "inline-labelled ABN")
    acn_inline = plant(doc, "AU_ACN", fmt_acn(ACN1), "detect", "inline-labelled ACN")
    # Canonical Luhn-valid test card (stable across Faker versions; Presidio-detectable).
    cc = plant(doc, "CREDIT_CARD", "4111 1111 1111 1111", "detect", "Luhn-valid card")
    iban = plant(doc, "IBAN_CODE", fake.iban(), "detect", "valid IBAN")
    org1 = plant(doc, "ORGANIZATION", "Halloway Supplies Pty Ltd", "detect", "vendor org name")
    vs.append([org1, f"ABN {abn_inline} / ACN {acn_inline}", cc, iban])
    # a BARE ABN in its own cell — no context word -> must be suppressed
    abn_bare = plant(doc, "AU_ABN", fmt_abn(ABN2), "suppress", "bare ABN, no context word")
    org2 = plant(doc, "ORGANIZATION", "Delcourt Freight Pty Ltd", "detect", "vendor org name")
    vs.append([org2, abn_bare, "", ""])

    # Sheet 3 — Journals: FP-TORTURE. No planted PII. 8-9 digit sequential IDs,
    # alnum codes, amounts, dates -> any detection here is an emergent false positive.
    js = wb.create_sheet("Journals")
    js.append(["JournalID", "Account", "Amount", "Effective Date", "GroupRef"])
    for i in range(220):
        # GroupRef: bare grouped digits (no phone/BSB context) — guards the phone
        # shape rule against flagging grouped account-style numbers.
        js.append([18400000 + i, f"EP{180000 + i}", round(1000.0 * (i - 110), 2),
                   f"2025-06-{(i % 28) + 1:02d}", f"{1000 + i} {8000 + i} {3000 + i}"])

    wb.save(str(path))


# --------------------------------------------------------------- PDF -----------

def build_pdf(path: Path) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception:
        return False
    doc = path.name
    styles = getSampleStyleSheet()
    d = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm,
                          topMargin=20 * mm, bottomMargin=20 * mm)

    med = complete_medicare("21234567")   # checksum-valid 10-digit Medicare
    partner = plant(doc, "PERSON", "Eleanor Whitcombe", "detect", "partner name in prose")
    email = plant(doc, "EMAIL_ADDRESS", "e.whitcombe@example.com", "detect", "email in prose")
    phone = plant(doc, "PHONE_NUMBER", "0412 345 678", "detect", "phone in prose w/ context")
    code = plant(doc, "ENGAGEMENT_CODE", "AUD-771204", "detect", "engagement code")
    addr = plant(doc, "AU_ADDRESS", "Unit 7, 14 Sturt Street, Ballarat VIC 3350", "detect",
                 "AU address in prose")
    tfn = plant(doc, "AU_TFN", fmt_tfn(TFN3), "detect", "inline-labelled TFN in prose")
    acn_bare = plant(doc, "AU_ACN", fmt_acn(complete_acn("41628054")), "suppress",
                     "bare ACN in prose, no label word")
    # now-covered types (recognizers added): in-context + valid values
    medicare = plant(doc, "MEDICARE", f"{med[:4]} {med[4:9]} {med[9:]}", "detect",
                     "checksum-valid Medicare, in-context")
    ip = plant(doc, "IP_ADDRESS", "203.0.113.47", "detect", "valid IP")
    dob = plant(doc, "DATE_OF_BIRTH", "14/03/1979", "detect", "DOB, in-context")
    passport = plant(doc, "AU_PASSPORT", "PA0123456", "detect", "passport, in-context")
    # one remaining documented gap — obsify has no SWIFT/BIC recognizer
    swift = plant(doc, "SWIFT_BIC", "CTBAAU2S", "gap", "no SWIFT/BIC recognizer")

    story = [
        Paragraph("Engagement Letter - Statutory Audit 2025", styles["Title"]),
        Spacer(1, 4 * mm),
        Paragraph(f"Led by partner {partner}, reachable at {email} or by phone on {phone} "
                  f"(reference {code}).", styles["BodyText"]),
        Paragraph(f"Please send documents to {addr}. Your nominated tax file number "
                  f"(TFN {tfn}) is on file.", styles["BodyText"]),
        Paragraph(f"The counterparty company number {acn_bare} was noted.", styles["BodyText"]),
        Paragraph(f"Medicare number {medicare}; access from {ip}; "
                  f"date of birth {dob}; passport {passport}.", styles["BodyText"]),
    ]
    d.build(story)
    return True


# --------------------------------------------------------------- DOCX ----------

def build_docx(path: Path) -> None:
    from docx import Document
    doc = path.name
    d = Document()
    d.add_heading("Audit Planning Memo", level=1)
    n1 = plant(doc, "PERSON", "Tomasz Kowalski", "detect", "name in paragraph")
    em = plant(doc, "EMAIL_ADDRESS", "t.kowalski@example.com", "detect", "email in paragraph")
    addr = plant(doc, "AU_ADDRESS", "Level 9, 320 Adelaide Street, Brisbane QLD 4000", "detect",
                 "AU address in paragraph")
    d.add_paragraph(f"Prepared by {n1} ({em}). Registered office: {addr}.")
    dl = plant(doc, "AU_DRIVER_LICENCE", "04938271", "detect", "driver licence, in-context")
    d.add_paragraph(f"Site contact driver licence {dl} verified on arrival.")

    d.add_heading("Vendors", level=2)
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Vendor", "Registration"
    abn = plant(doc, "AU_ABN", fmt_abn(_mint_abn(72155399000)), "detect",
                "inline-labelled ABN in table cell")
    org = plant(doc, "ORGANIZATION", "Marrick Timber Pty Ltd", "detect", "vendor org name")
    r = t.add_row().cells
    r[0].text, r[1].text = org, f"ABN {abn}"
    d.save(str(path))


# --------------------------------------------------------------- driver --------

def build_eval_corpus(out_dir: str) -> dict:
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    _KEY.clear()
    Faker.seed(7)          # reseed so every call is byte-for-byte reproducible
    files = []

    build_excel(out / "eval_ledger.xlsx"); files.append("eval_ledger.xlsx")
    build_docx(out / "eval_memo.docx"); files.append("eval_memo.docx")
    notes = []
    if build_pdf(out / "eval_engagement.pdf"):
        files.append("eval_engagement.pdf")
    else:
        notes.append("reportlab missing; skipped eval_engagement.pdf (PDF-only planted items absent)")

    key_path = out / "answer_key.jsonl"
    with key_path.open("w", encoding="utf-8") as fh:
        for item in _KEY:
            fh.write(json.dumps(item) + "\n")

    return {"files": files, "answer_key": str(key_path), "planted": len(_KEY), "notes": notes}


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(prog="python eval/generate.py",
                                 description="Generate a labelled synthetic eval corpus (all fake).")
    ap.add_argument("--out", default="./eval_corpus")
    args = ap.parse_args(argv)
    r = build_eval_corpus(args.out)
    print(f"[generate] wrote {len(r['files'])} document(s) + answer key ({r['planted']} planted items) to {args.out}")
    for f in r["files"]:
        print(f"  - {f}")
    for n in r["notes"]:
        print(f"[generate] note: {n}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
